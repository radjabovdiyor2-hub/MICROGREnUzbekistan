"""
Сквозная сверка пакета: Word ↔ Excel ↔ model.py.

Пакет разъехался потому, что одни и те же величины набирались в четырёх местах
руками, а проверить их было нечем. Эта сверка читает готовые документы и
сравнивает с моделью — то есть проверяет РЕЗУЛЬТАТ сборки, а не намерение.

Запуск:
    C:/Users/noteb/AppData/Local/Programs/Python/Python310/python.exe -X utf8 \
        doc/mega_audit.py

Сверка обязана уметь краснеть. Подсуньте заведомо неверную цифру в любой из
четырёх блоков — падать должен именно он.
"""

from __future__ import annotations

import re
import sys
from datetime import date
from pathlib import Path

import docx
import openpyxl

sys.path.insert(0, str(Path(__file__).resolve().parent))
import model as M  # noqa: E402

HERE = Path(__file__).resolve().parent
READY = HERE / "ready"

NAMES = {
    "ariza": "Ariza_Microgreen_Uzbekistan.docx",
    "biznes": "Biznes_reja_Microgreen_Uzbekistan.docx",
    "sotuv": "Sotuv_reja_Microgreen_Uzbekistan.docx",
    "kpi": "KPI_Microgreen_Uzbekistan.docx",
    "chora": "Chora_tadbirlar_Microgreen_Uzbekistan.docx",
    "excel": "Fin_Model_Microgreen_Uzbekistan.xlsx",
}

problems: list[str] = []
checks = 0


def fail(block: str, message: str) -> None:
    problems.append(f"[{block}] {message}")


def ok(block: str, message: str) -> None:
    global checks
    checks += 1
    print(f"  ok  {block}: {message}")


def doc_text(key: str) -> str:
    document = docx.Document(READY / NAMES[key])
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def numbers(text: str) -> set[int]:
    """Все целые числа текста, без разделителей."""
    found = set()
    for raw in re.findall(r"\d[\d \u00a0.,]*\d|\d", text):
        digits = re.sub(r"[^\d]", "", raw)
        if digits:
            found.add(int(digits))
    return found


# ══════════════════════════════════════════════════════════════════════════
# 1. Бюджет
# ══════════════════════════════════════════════════════════════════════════

def check_budget(m: dict) -> None:
    block = "byudjet"
    items = sum(M.BUDGET_USD.values())

    if items + M.RESERVE_USD != M.INVESTMENT_USD:
        fail(block, f"moddalar {items} + zaxira {M.RESERVE_USD} != "
                    f"so'ralayotgan {M.INVESTMENT_USD}")
    else:
        ok(block, f"${items} + ${M.RESERVE_USD} zaxira = ${M.INVESTMENT_USD}")

    chora = doc_text("chora")
    chora_numbers = numbers(chora)
    for name, usd in M.BUDGET_USD.items():
        if usd not in chora_numbers:
            fail(block, f"Chora-tadbirlar: moddasi «{name}» = ${usd} topilmadi")
    if M.INVESTMENT_USD not in chora_numbers:
        fail(block, f"Chora-tadbirlar: jami ${M.INVESTMENT_USD} yo'q")
    else:
        ok(block, "Chora-tadbirlar: barcha moddalar va jami mos")

    biznes = doc_text("biznes")
    biznes_numbers = numbers(biznes)
    for name, usd in M.BUDGET_USD.items():
        if usd not in biznes_numbers:
            fail(block, f"Biznes-reja: moddasi «{name}» = ${usd} topilmadi")
    if M.RESERVE_USD not in biznes_numbers:
        fail(block, f"Biznes-reja: zaxira ${M.RESERVE_USD} alohida ko'rsatilmagan")
    else:
        ok(block, "Biznes-reja: byudjet Chora-tadbirlar bilan bir xil")

    # CAPEX в Excel обязан равняться капитальным статьям бюджета.
    wb = openpyxl.load_workbook(READY / NAMES["excel"])
    ws = wb["5. Aylanma mablag'"]
    capex_row = next((r for r in ws.iter_rows()
                      if r[0].value and "CAPEX" in str(r[0].value)), None)
    if capex_row is None:
        fail(block, "Excel: CAPEX qatori topilmadi")
        return
    capex_excel = sum(c.value for c in capex_row[1:M.MONTHS + 1] if isinstance(c.value, int))
    if capex_excel != m["capex_uzs"]:
        fail(block, f"Excel CAPEX {capex_excel:,} != byudjet kapital moddalari "
                    f"{m['capex_uzs']:,}")
    else:
        ok(block, f"CAPEX {capex_excel:,} = kapital moddalar summasi")


# ══════════════════════════════════════════════════════════════════════════
# 2. Календарь и стартовая точка
# ══════════════════════════════════════════════════════════════════════════

def check_calendar(m: dict) -> None:
    block = "kalendar"
    today = date.today()
    start = date(M.START_YEAR, M.START_MONTH, 1)
    if start < date(today.year, today.month, 1):
        fail(block, f"model {m['labels'][0]} dan boshlanadi — bu o'tgan davr "
                    f"({today:%Y-%m}). Hech bir oy o'tmishda bo'lmasligi kerak")
    else:
        ok(block, f"davr {m['labels'][0]} — {m['labels'][-1]}, o'tmishda oy yo'q")

    tranche_month = next((i for i, v in enumerate(m["investment"]) if v), None)
    if tranche_month != 0:
        fail(block, f"transh M{(tranche_month or 0) + 1} da — birinchi oyda bo'lishi kerak")
    else:
        ok(block, f"transh {m['labels'][0]} da, ya'ni birinchi oyda")

    # Заявка и модель обязаны сходиться в стартовой точке.
    ariza = doc_text("ariza")
    base_mln = round(m["revenue"][0] / 1_000_000)
    if str(base_mln) not in ariza:
        fail(block, f"Ariza 32-bandida bazaviy daromad {base_mln} mln topilmadi")
    elif str(m["restaurants"][0]) not in ariza or str(m["families"][0]) not in ariza:
        fail(block, f"Ariza: baza {m['restaurants'][0]} restoran + "
                    f"{m['families'][0]} oila ko'rsatilmagan")
    else:
        ok(block, f"Ariza 32-band = model M1: {base_mln} mln, "
                  f"{m['restaurants'][0]} restoran + {m['families'][0]} oila")

    # Сумма прежних вложений основателя — главный аргумент заявки. Без неё
    # фонд не видит, что у основателя своя шкура в игре.
    if str(M.FOUNDER_INVESTED_USD // 1000) not in ariza:
        fail(block, f"Ariza 34-bandida asoschi kiritgan "
                    f"${M.FOUNDER_INVESTED_USD:,} ko'rsatilmagan")
    else:
        ok(block, f"Ariza 34-band: asoschi ${M.FOUNDER_INVESTED_USD:,} kiritgan")


# ══════════════════════════════════════════════════════════════════════════
# 2b. Каналы продаж: три потока обязаны складываться в выручку
# ══════════════════════════════════════════════════════════════════════════

def check_channels(m: dict) -> None:
    block = "kanallar"

    for i in range(M.MONTHS):
        total = (m["b2b"][i] + m["subscriptions"][i]
                 + m["retail"][i] + m["strawberry"][i])
        if total != m["revenue"][i]:
            fail(block, f"{m['labels'][i]}: kanallar yig'indisi {total:,} != "
                        f"daromad {m['revenue'][i]:,}")
            return
    ok(block, f"barcha {M.MONTHS} oyda uch kanal + qulupnay = jami daromad")

    # ARPU ресторана обязан быть произведением чека на частоту, а не числом
    # «из головы»: прежние документы писали 200 000 как месячный ARPU, из-за
    # чего выручка в таблицах не сходилась с выручкой в соседней строке.
    expected = M.CHECK_RESTAURANT * M.DELIVERIES_PER_MONTH
    if m["arpu_restaurant"] != expected:
        fail(block, f"ARPU restoran {m['arpu_restaurant']:,} != chek × chastota "
                    f"{expected:,}")
    else:
        ok(block, f"ARPU restoran = {M.CHECK_RESTAURANT:,} × "
                  f"{M.DELIVERIES_PER_MONTH} = {expected:,}")

    # Доля ресторанов — то, что назвал владелец: около 30 %.
    share = m["b2b"][0] / m["revenue"][0]
    if not 0.25 <= share <= 0.35:
        fail(block, f"restoranlar ulushi {share:.0%} — egasining "
                    f"«~30%» ma'lumotiga mos emas")
    else:
        ok(block, f"M1 tarkibi: restoranlar {share:.0%}, qolgani "
                  f"{1 - share:.0%} (obuna + chakana)")


# ══════════════════════════════════════════════════════════════════════════
# 3. Траектория: квартальные таблицы против модели
# ══════════════════════════════════════════════════════════════════════════

def check_trajectory(m: dict) -> None:
    block = "trayektoriya"
    quarters = M.quarters(m)[:4]

    biznes = docx.Document(READY / NAMES["biznes"])
    plan = biznes.tables[6]

    def row_values(needle: str) -> list[str] | None:
        for row in plan.rows:
            if needle.lower() in row.cells[0].text.strip().lower():
                return [c.text.strip() for c in row.cells[1:5]]
        return None

    customers = row_values("Jami mijozlar")
    if customers is None:
        fail(block, "Biznes-reja: «Jami mijozlar» qatori yo'q")
    else:
        expected = [str(q["customers"]) for q in quarters]
        if customers != expected:
            fail(block, f"Biznes-reja mijozlar {customers} != model {expected}")
        else:
            ok(block, f"Biznes-reja choraklik mijozlar = model {expected}")

    restaurants = row_values("Restoranlar")
    if restaurants is not None:
        expected = [str(q["restaurants"]) for q in quarters]
        if restaurants != expected:
            fail(block, f"Biznes-reja restoranlar {restaurants} != model {expected}")
        else:
            ok(block, "Biznes-reja choraklik restoranlar = model")

    # Битая ячейка, которая жила в пакете: в строке Marketing стояло слово.
    marketing = row_values("Marketing")
    if marketing and any(not re.match(r"^[\d.,\s—-]*$", v) for v in marketing):
        fail(block, f"Biznes-reja: «Marketing» qatorida raqam emas: {marketing}")
    else:
        ok(block, "Biznes-reja: «Marketing» qatorida faqat raqamlar")

    # Sotuv-reja обязан говорить то же самое.
    sotuv = docx.Document(READY / NAMES["sotuv"])
    sotuv_table = sotuv.tables[0]
    for i, q in enumerate(quarters):
        if i + 1 >= len(sotuv_table.rows):
            break
        cell = sotuv_table.rows[i + 1].cells[1].text
        if str(q["customers"]) not in cell:
            fail(block, f"Sotuv-reja {i + 1}-chorak: «{cell}» != "
                        f"model {q['customers']} mijoz")
    else:
        ok(block, "Sotuv-reja choraklari = model")


# ══════════════════════════════════════════════════════════════════════════
# 4. Юнит-экономика: одна версия во всех документах
# ══════════════════════════════════════════════════════════════════════════

def check_unit_economics(m: dict) -> None:
    block = "unit-ekonomika"
    last = M.MONTHS - 1

    biznes = docx.Document(READY / NAMES["biznes"])
    unit = biznes.tables[2]
    text = "\n".join(c.text for r in unit.rows for c in r.cells)
    found = numbers(text)

    if m["arpu_restaurant"] not in found:
        fail(block, f"Biznes-reja: ARPU restoran {m['arpu_restaurant']:,} yo'q "
                    f"(chek × chastota bo'lishi shart)")
    elif M.PRICE_FAMILY not in found:
        fail(block, f"Biznes-reja: oila obunasi {M.PRICE_FAMILY:,} yo'q")
    elif M.CHECK_RETAIL not in found:
        fail(block, f"Biznes-reja: chakana chek {M.CHECK_RETAIL:,} yo'q")
    else:
        ok(block, f"ARPU restoran {m['arpu_restaurant']:,} / oila "
                  f"{M.PRICE_FAMILY:,} / chakana {M.CHECK_RETAIL:,}")

    ratio = m["ltv"][last] / m["cac"][last]
    if f"{ratio:.0f}x" not in text:
        fail(block, f"Biznes-reja: LTV/CAC {ratio:.0f}x topilmadi")
    else:
        ok(block, f"LTV/CAC {ratio:.0f}x = model")

    if "ARPU × brutto marja" not in text:
        fail(block, "Biznes-reja: LTV formulasi ko'rsatilmagan — "
                    "raqamga ishonish uchun formula kerak")
    else:
        ok(block, "LTV formulasi hujjatda ko'rsatilgan")

    sotuv = doc_text("sotuv")
    if f"{M.CHECK_RESTAURANT // 1000}k" not in sotuv:
        fail(block, "Sotuv-reja: ARPU Biznes-reja bilan mos emas")
    else:
        ok(block, "Sotuv-reja unit-ekonomikasi = Biznes-reja")


# ══════════════════════════════════════════════════════════════════════════
# 5. KPI: каждый порог НИЖЕ модели
# ══════════════════════════════════════════════════════════════════════════

def check_kpi(m: dict) -> None:
    block = "kpi"
    kpi = docx.Document(READY / NAMES["kpi"])

    def target(table_index: int, needle: str) -> tuple[str, int] | None:
        for row in kpi.tables[table_index].rows:
            if needle.lower() in row.cells[1].text.strip().lower():
                months = re.search(r"(\d+)\s*oy", row.cells[3].text)
                return row.cells[2].text.strip(), int(months.group(1)) if months else 0
        return None

    cases = [
        (0, "Yangi mijozlar", "restoran", "restaurants"),
        (1, "Restoranlar", "restoran", "restaurants"),
        (1, "Oilalar", "oila", "families"),
    ]
    for table_index, needle, _unit, series in cases:
        found = target(table_index, needle)
        if not found:
            fail(block, f"«{needle}» qatori topilmadi")
            continue
        text, months = found
        promised = re.search(r"(\d+)\+", text)
        if not promised or not months:
            fail(block, f"«{needle}»: maqsad yoki muddat o'qilmadi ({text})")
            continue
        promised_value = int(promised.group(1))
        actual = m[series][months - 1]
        today_base = m[series][0]
        if promised_value >= actual:
            fail(block, f"«{needle}»: KPI {promised_value}+ >= model {actual} "
                        f"({months} oy) — zaxira yo'q, majburiyat bajarilmaydi")
        elif promised_value <= today_base:
            # Порог ниже сегодняшней базы = обязательство, выполненное в момент
            # подписания. Фонд платит транш за результат, которого уже достигли.
            fail(block, f"«{needle}»: KPI {promised_value}+ <= bugungi baza "
                        f"{today_base} — imzolash paytidayoq bajarilgan bo'ladi")
        else:
            ok(block, f"«{needle}» {today_base} (bugun) < {promised_value}+ "
                      f"< {actual} (model, {months} oy)")

    for table_index in (0, 1):
        found = target(table_index, "Oylik daromad")
        if not found:
            continue
        text, months = found
        promised = re.search(r"(\d+)\s*mln", text)
        if not promised or not months:
            continue
        promised_value = int(promised.group(1)) * 1_000_000
        actual = m["revenue"][months - 1]
        if promised_value >= actual:
            fail(block, f"«Oylik daromad»: KPI {promised_value:,} >= model "
                        f"{actual:,} ({months} oy)")
        else:
            ok(block, f"«Oylik daromad» {promised_value // 10**6} mln < model "
                      f"{actual // 10**6} mln ({months} oy)")


# ══════════════════════════════════════════════════════════════════════════

def main() -> int:
    m = M.build()
    print("СКВОЗНАЯ СВЕРКА ПАКЕТА\n")
    for name in NAMES.values():
        if not (READY / name).exists():
            print(f"[FAIL] нет файла: {name}")
            return 1

    check_budget(m)
    check_calendar(m)
    check_channels(m)
    check_trajectory(m)
    check_unit_economics(m)
    check_kpi(m)

    print(f"\nпроверок пройдено: {checks}")
    if problems:
        print(f"\n[FAIL] расхождений: {len(problems)}")
        for item in problems:
            print(f"  · {item}")
        return 1
    print("\n[OK] пакет согласован: Word ↔ Excel ↔ модель")
    return 0


if __name__ == "__main__":
    sys.exit(main())
