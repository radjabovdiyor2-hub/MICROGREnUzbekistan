"""
Сборка пяти Word-документов пакета из модели.

Документы не создаются с нуля: формы `Ariza` и `KPI` — бланки фонда, их
структуру ломать нельзя. Берём оригиналы из `backup-2026-08-12/`, подставляем
цифры из `model.py` и сохраняем в `ready/`. Форматирование сохраняется, потому
что правится текст существующих runs, а не пересобирается абзац.

Запуск:
    C:/Users/noteb/AppData/Local/Programs/Python/Python310/python.exe -X utf8 \
        doc/build_docs.py
"""

from __future__ import annotations

import sys
from datetime import date
from pathlib import Path

import docx
from copy import deepcopy

from docx.oxml.ns import qn

sys.path.insert(0, str(Path(__file__).resolve().parent))
import model as M  # noqa: E402

HERE = Path(__file__).resolve().parent
SRC = HERE / "backup-2026-08-12"
OUT = HERE / "ready"

NAMES = {
    "ariza": "Ariza_Microgreen_Uzbekistan.docx",
    "biznes": "Biznes_reja_Microgreen_Uzbekistan.docx",
    "sotuv": "Sotuv_reja_Microgreen_Uzbekistan.docx",
    "kpi": "KPI_Microgreen_Uzbekistan.docx",
    "chora": "Chora_tadbirlar_Microgreen_Uzbekistan.docx",
}


def mln(value: int) -> str:
    """12 345 678 → «12,3 mln». Одинаково во всех документах."""
    return f"{value / 1_000_000:.1f}".replace(".0", "") + " mln"


def set_cell(cell, text: str) -> None:
    """Заменить текст ячейки, сохранив оформление первого run."""
    paragraphs = cell.paragraphs
    if paragraphs and paragraphs[0].runs:
        paragraphs[0].runs[0].text = text
        for run in paragraphs[0].runs[1:]:
            run.text = ""
        for extra in paragraphs[1:]:
            for run in extra.runs:
                run.text = ""
    else:
        cell.text = text


def set_para(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = text


def unmerge_first_cell(row) -> None:
    """Разъединить подпись строки, если она растянута на несколько колонок.

    В бланке строка «Marketing» собрана с gridSpan=2: подпись занимает две
    колонки, и на кварталы остаётся три ячейки вместо четырёх. Именно поэтому
    в исходном пакете эта строка выглядела как «Marketing | Marketing | 1.5 |
    1.0 | 1.0» — один квартал просто не поместился. Возвращаем строке обычную
    сетку, чтобы она встала в один ряд с остальными.
    """
    tc = row.cells[0]._tc
    tcPr = tc.tcPr
    if tcPr is None:
        return
    grid_span = tcPr.find(qn("w:gridSpan"))
    if grid_span is None:
        return
    span = int(grid_span.get(qn("w:val")))
    tcPr.remove(grid_span)
    anchor = tc
    for _ in range(span - 1):
        clone = deepcopy(tc)
        for paragraph in clone.findall(qn("w:p")):
            for run in paragraph.findall(qn("w:r")):
                paragraph.remove(run)
        anchor.addnext(clone)
        anchor = clone


def find_row(table, needle: str) -> int | None:
    for i, row in enumerate(table.rows):
        if needle.lower() in row.cells[0].text.strip().lower():
            return i
    return None


# ══════════════════════════════════════════════════════════════════════════
# KPI-пороги: считаются ОТ модели и всегда ниже её значений.
#
# В прежней версии было наоборот: «10+ новых ресторанов (всего 15+) за 4
# месяца», тогда как модель на ту же дату давала 14 и +5. Подписать такой KPI —
# значит обязаться к недостижимому.
# ══════════════════════════════════════════════════════════════════════════

def kpi_targets(m: dict) -> tuple[list[tuple], list[tuple]]:
    strawberry_kg = [round(v / M.STRAWBERRY_PRICE_KG) for v in m["strawberry"]]
    kg_by_m4 = sum(strawberry_kg[:4])

    tranche1 = [
        ("Qulupnay ferma qurish",
         f"60 m2, {M.STRAWBERRY_PLANTS // 180} premium stelaj, NFT gidroponik tizim",
         "3 oy"),
        ("Birinchi hosil",
         f"{int(kg_by_m4 * 0.8 // 10) * 10}+ kg qulupnay yig'ilgan va sotilgan",
         "4 oy"),
        ("Yangi mijozlar",
         f"{m['restaurants'][4] - 1}+ restoran (jami), {m['families'][4] - 3}+ oila",
         "5 oy"),
        ("Oylik daromad", f"{int(m['revenue'][5] * 0.85 / 1e6)} mln so'mdan oshishi", "6 oy"),
        ("IT-platforma", "100+ foydalanuvchi saytda va botda", "6 oy"),
    ]
    tranche2 = [
        ("Oylik daromad", f"{int(m['revenue'][7] * 0.85 / 1e6)} mln so'mdan ko'p", "8 oy"),
        ("Restoranlar", f"{m['restaurants'][9] - 1}+ barqaror restoran", "10 oy"),
        ("Oilalar", f"{m['families'][9] - 3}+ uy xo'jaligi", "10 oy"),
        ("Toshkent", "2+ Toshkent restorani bilan shartnoma", "11 oy"),
        ("Xodimlar", "5+ kishi doimiy ish o'rni", "12 oy"),
        ("Rentabellik",
         f"Ferma to'liq quvvatda, oylik sof foyda {int(m['net'][11] * 0.8 / 1e6)} mln+",
         "12 oy"),
    ]
    return tranche1, tranche2


# ══════════════════════════════════════════════════════════════════════════

def build_chora(m: dict) -> None:
    """Бюджет: статьи $30 000 + резерв $5 000 = $35 000, резерв один раз."""
    doc = docx.Document(SRC / NAMES["chora"])
    table = doc.tables[0]

    rows = [
        ("Bino ijarasi va tayyorlash",
         "60 m2, ijara 6 oy oldindan, ta'mirlash, elektr 15kVt, suv", "1 oy",
         "Bino tayyor"),
        ("Premium gidroponik uskunalar",
         "15 stelaj, 45 NFT lotok, 2700 gorshok, nasoslar", "1-2 oy",
         "Vertikal ferma o'rnatilgan"),
        ("LED + klimat",
         "Fito-lampalar, konditsioner, ventilyatsiya, nazorat", "2-3 oy",
         "Mikroiqlim ta'minlangan"),
        ("Ekin materiallari",
         f"{M.STRAWBERRY_PLANTS} ko'chat, substrat, ishga tushirish", "3-4 oy",
         "Birinchi hosil olindi"),
        ("Marketing + IT + aylanma",
         "Brend, reklama, degustatsiya, IT-platforma, aylanma mablag'", "1-12 oy",
         "30+ restoran, 130+ oila"),
    ]
    keys = list(M.BUDGET_USD)

    # В бланке было четыре статьи, стало пять: «LED + klimat» и «Ekin
    # materiallari» разделены, потому что первое — оборудование (амортизируется),
    # второе — расходный материал. Итоговая строка одна, её держим последней.
    needed = len(rows) + 2                       # шапка + статьи + ЖАМИ
    while len(table.rows) < needed:
        table.add_row()

    for i, (name, how, when, result) in enumerate(rows):
        row = table.rows[i + 1]
        usd = M.BUDGET_USD[keys[i]]
        set_cell(row.cells[0], f"{i + 1}.")
        set_cell(row.cells[1], name)
        set_cell(row.cells[2], how)
        set_cell(row.cells[3], when)
        set_cell(row.cells[4], f"${usd:,} ({mln(usd * M.USD_RATE)} so'm)".replace(",", " "))
        set_cell(row.cells[5], result)

    total_row = table.rows[len(rows) + 1]
    items = sum(M.BUDGET_USD.values())
    set_cell(total_row.cells[0], "")
    set_cell(total_row.cells[1], "JAMI moddalar")
    set_cell(total_row.cells[2], f"Zaxira alohida: ${M.RESERVE_USD:,}".replace(",", " "))
    set_cell(total_row.cells[3], "1-12 oy")
    set_cell(total_row.cells[4],
             f"${items:,} + ${M.RESERVE_USD:,} zaxira = ${M.INVESTMENT_USD:,}"
             .replace(",", " "))
    set_cell(total_row.cells[5],
             f"Jami {mln(M.INVESTMENT_UZS)} so'm")

    doc.save(OUT / NAMES["chora"])


def build_kpi(m: dict) -> None:
    t1, t2 = kpi_targets(m)
    doc = docx.Document(SRC / NAMES["kpi"])

    for table, targets in ((doc.tables[0], t1), (doc.tables[1], t2)):
        data_rows = [r for r in table.rows[1:] if r.cells[0].text.strip()
                     or r.cells[1].text.strip()]
        for i, row in enumerate(data_rows):
            if i >= len(targets):
                for cell in row.cells:
                    set_cell(cell, "")
                continue
            name, requirement, term = targets[i]
            set_cell(row.cells[0], str(i + 1))
            set_cell(row.cells[1], name)
            set_cell(row.cells[2], requirement)
            set_cell(row.cells[3], term)

    doc.save(OUT / NAMES["kpi"])


def build_sotuv(m: dict) -> None:
    quarters = M.quarters(m)
    doc = docx.Document(SRC / NAMES["sotuv"])

    kg_month = round(M.STRAWBERRY_PLANTS * M.KG_PER_PLANT_YEAR / 12)
    last = len(m["labels"]) - 1

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("1. Maqsadlar"):
            # Здесь стояло «720 kg/oy» и «20 → 200+ mijoz, 200 mln» — цифры,
            # которых нет ни в модели, ни в физике фермы 60 м².
            set_para(paragraph,
                     f"1. Maqsadlar\n"
                     f"• Samarqand va Toshkentda mikro-ko'katlar + qulupnay "
                     f"yetkazuvchi №1\n"
                     f"• Premium qulupnay siti-fermani ishga tushirish "
                     f"({kg_month} kg/oy — 60 m² quvvati)\n"
                     f"• Mijozlar: {m['customers'][0]} → {m['customers'][last]}, "
                     f"daromad: {mln(m['revenue'][0])} → {mln(m['revenue'][last])} so'm")
        elif text.startswith("2. Traction"):
            set_para(paragraph,
                     f"2. Traction (bugungi holat)\n"
                     f"• {m['restaurants'][0]} restoran + {m['families'][0]} oila = "
                     f"{m['customers'][0]} barqaror mijoz\n"
                     f"• Oylik: {mln(m['revenue'][0])} so'm | microgreenuzbekistan.com\n"
                     f"• \"Mirzo Ulug'bek Vorislari\" — Samarqand viloyat g'olibi\n"
                     f"• Kompaniya bugun ham foydali: sof foyda "
                     f"{mln(m['net'][0])} so'm/oy")
        elif text.startswith("3. Ijro rejasi"):
            lines = ["3. Ijro rejasi (investitsiyadan keyin)"]
            for q in quarters[:4]:
                lines.append(f"{q['label']}: {q['restaurants']} rest, "
                             f"{q['families']} oila — {mln(q['revenue_month'])} so'm/oy")
            set_para(paragraph, "\n".join(lines))
        elif text.startswith("7. Unit Economics"):
            set_para(paragraph,
                     f"7. Unit Economics: ARPU {mln(M.PRICE_RESTAURANT)} (restoran) / "
                     f"{M.PRICE_FAMILY // 1000}k (oila) | "
                     f"CAC ~{m['cac'][last] // 1000}k | "
                     f"LTV/CAC {m['ltv'][last] / m['cac'][last]:.0f}x | "
                     f"Brutto marja {m['gross_margin'][last] * 100:.0f}%")
        elif text.startswith("4. Rivojlanish"):
            set_para(paragraph,
                     f"4. Rivojlanish: Qulupnay ferma ({quarters[0]['label']}) | "
                     f"to'liq quvvat ({m['labels'][6]}) | Toshkent bozori | "
                     f"Supermarketlar")

    table = doc.tables[0]
    for i, q in enumerate(quarters[:4]):
        row = table.rows[i + 1]
        set_cell(row.cells[0], q["label"])
        set_cell(row.cells[1], f"{q['customers']} mijoz")
        set_cell(row.cells[2], mln(q["revenue_month"]))
        set_cell(row.cells[3], ["Ferma qurilishi", "Birinchi hosil",
                                "To'liq quvvat", "Toshkent bozori"][i])
    set_cell(table.rows[0].cells[1], "Jami mijozlar")
    set_cell(table.rows[0].cells[2], "Oylik daromad")

    doc.save(OUT / NAMES["sotuv"])


def build_biznes(m: dict) -> None:
    quarters = M.quarters(m)
    doc = docx.Document(SRC / NAMES["biznes"])
    last = len(m["labels"]) - 1

    # Таблица 2 — техническая справка. В системе 12 ботов + мост, не 11.
    tech = doc.tables[1]
    row_index = find_row(tech, "7")
    for row in tech.rows:
        if "AI-ofis" in row.cells[1].text:
            set_cell(row.cells[2], "12 ta AI-bot + n8n ko'prigi (CRM, sotuv, moliya)")

    # Таблица 3 — юнит-экономика. Одна версия, из модели, с формулами.
    unit = doc.tables[2]
    values = [
        ("ARPU (oy)",
         f"{M.PRICE_RESTAURANT:,} so'm (restoran), {M.PRICE_FAMILY:,} (oila)"
         .replace(",", " "),
         "Tariflar varag'idagi obuna narxlari"),
        ("CAC", f"~{m['cac'][last]:,} so'm".replace(",", " "),
         "(Marketing + direktor ish haqining 40%) / yangi mijozlar"),
        ("LTV", f"{m['ltv'][last]:,} so'm".replace(",", " "),
         f"ARPU × brutto marja × min(1/churn, {M.MONTHS} oy)"),
        ("LTV/CAC", f"{m['ltv'][last] / m['cac'][last]:.0f}x", "(LTV / CAC)"),
        ("Payback (oy)",
         f"{m['cac'][last] / (m['arpu'][last] * m['gross_margin'][last]):.1f} oy",
         "(CAC / (ARPU × brutto marja))"),
        ("Churn (oy)", f"{M.MONTHLY_CHURN * 100:.1f}%",
         "(Yo'qotilgan mijozlar / umumiy mijozlar)"),
        ("Brutto marja", f"{m['gross_margin'][last] * 100:.0f}%",
         "Obuna 70%, qulupnay 55% — aralash"),
    ]
    for i, (name, value, formula) in enumerate(values):
        if i + 1 >= len(unit.rows):
            break
        row = unit.rows[i + 1]
        set_cell(row.cells[0], name)
        set_cell(row.cells[1], value)
        set_cell(row.cells[2], formula)

    # Таблица 6 — бюджет, тот же, что в Chora-tadbirlar.
    budget = doc.tables[5]
    keys = list(M.BUDGET_USD)
    notes = ["Ijara 6 oy oldindan + ta'mirlash", "15 stelaj + NFT lotoklar",
             "Fito-lampalar + konditsioner", f"{M.STRAWBERRY_PLANTS} ko'chat + substrat",
             "Brend, reklama, IT, aylanma mablag'"]
    for i, key in enumerate(keys):
        if i + 1 >= len(budget.rows):
            break
        usd = M.BUDGET_USD[key]
        row = budget.rows[i + 1]
        set_cell(row.cells[0], key)
        set_cell(row.cells[1], f"${usd:,}".replace(",", " "))
        set_cell(row.cells[2], f"{usd / M.INVESTMENT_USD * 100:.1f}%")
        set_cell(row.cells[3], notes[i])

    # Резерв показываем отдельной строкой — раньше он сидел внутри маркетинга
    # и одновременно в итоге, из-за чего сумма не сходилась.
    extra = budget.add_row()
    set_cell(extra.cells[0], "Zaxira (kutilmagan xarajatlar)")
    set_cell(extra.cells[1], f"${M.RESERVE_USD:,}".replace(",", " "))
    set_cell(extra.cells[2], f"{M.RESERVE_USD / M.INVESTMENT_USD * 100:.1f}%")
    set_cell(extra.cells[3], "Alohida saqlanadi, moddalarga kirmaydi")

    total = budget.add_row()
    set_cell(total.cells[0], "JAMI")
    set_cell(total.cells[1], f"${M.INVESTMENT_USD:,}".replace(",", " "))
    set_cell(total.cells[2], "100.0%")
    set_cell(total.cells[3], f"{mln(M.INVESTMENT_UZS)} so'm")

    # Таблица 7 — квартальный план, собирается агрегацией модели.
    plan = doc.tables[6]
    quarter_labels = [q["label"] for q in quarters[:4]]

    def distinct(row):
        """Ячейки без повторов.

        В бланке часть строк собрана с объединением: python-docx отдаёт одну и
        ту же ячейку дважды. Из-за этого запись значения в «вторую» ячейку
        затирала подпись строки — именно так в исходном пакете строка
        «Marketing» лишилась названия и осталась с числом вместо него.
        """
        result = []
        for cell in row.cells:
            if not result or cell._tc is not result[-1]._tc:
                result.append(cell)
        return result

    def fill(needle: str, values: list[str]) -> None:
        idx = find_row(plan, needle)
        if idx is None:
            return
        row = plan.rows[idx]
        if len(distinct(row)) < len(values) + 1:
            unmerge_first_cell(row)
        cells = distinct(row)
        set_cell(cells[0], needle)
        for col, value in enumerate(values):
            target = col + 1
            if target < len(cells):
                set_cell(cells[target], value)

    fill("Year", ["2026-27"] * 4)
    fill("Quarter", quarter_labels)
    fill("Jami mijozlar", [str(q["customers"]) for q in quarters[:4]])
    fill("Yangi mijozlar", [str(q["new_customers"]) for q in quarters[:4]])
    fill("Ketgan mijozlar", ["—"] * 4)
    fill("Restoranlar", [str(q["restaurants"]) for q in quarters[:4]])
    fill("Oilalar", [str(q["families"]) for q in quarters[:4]])
    fill("Premium mijozlar", [str(q["restaurants"]) for q in quarters[:4]])
    fill("Valyuta kursi", [f"{M.USD_RATE:,}".replace(",", " ")] * 4)
    fill("Chek restoran", [f"{M.PRICE_RESTAURANT:,}".replace(",", " ")] * 4)
    fill("Chek oila", [f"{M.PRICE_FAMILY:,}".replace(",", " ")] * 4)
    fill("Mikro-ko'katlar",
         [f"{(m['b2b'][i * 3 + 2] + m['b2c'][i * 3 + 2]) / 1e6:.1f}" for i in range(4)])
    fill("Qulupnay", [f"{q['strawberry_month'] / 1e6:.1f}" for q in quarters[:4]])

    for key in ("Ish haqi", "Ijara+kommunal", "Marketing", "Elektr energiya",
                "Materiallar", "Soliq"):
        source = {"Ijara+kommunal": "Ijara + kommunal", "Soliq": "Soliq (4%)"}.get(key, key)
        if source in m["opex_rows"]:
            fill(key, [f"{m['opex_rows'][source][i * 3 + 2] / 1e6:.1f}" for i in range(4)])

    for key in ("Bank xizmati", "Komissiyalar"):
        idx = find_row(plan, key)
        if idx is not None:
            for cell in plan.rows[idx].cells:
                set_cell(cell, "")

    fill("Brutto foyda", [f"{q['gross_sum'] / 3 / 1e6:.1f}" for q in quarters[:4]])
    fill("EBITDA",
         [f"{sum(m['ebitda'][i * 3:(i + 1) * 3]) / 3 / 1e6:.1f}" for i in range(4)])
    fill("Sof foyda", [f"{q['net_sum'] / 3 / 1e6:.1f}" for q in quarters[:4]])
    fill("Kassada pul", [f"{q['cash_end'] / 1e6:.0f}" for q in quarters[:4]])
    fill("Break-even", ["O'tilgan"] * 4)

    doc.save(OUT / NAMES["biznes"])


def build_ariza(m: dict) -> None:
    quarters = M.quarters(m)
    doc = docx.Document(SRC / NAMES["ariza"])
    table = doc.tables[1]

    plan = "; ".join(
        f"{q['label']}: {mln(q['revenue_month'])} so'm/oy "
        f"({q['restaurants']} restoran + {q['families']} oila)"
        for q in quarters[:4]
    )
    # Возраст считается на дату сборки пакета, а не вписывается руками: в
    # исходном документе стояло «23 yosh», и оно становилось неверным через
    # несколько дней (день рождения 24 августа).
    today = date.today()
    born = date(2002, 8, 24)
    age = today.year - born.year - ((today.month, today.day) < (born.month, born.day))

    last = len(m["labels"]) - 1
    updates = {
        "7": f"$300 000 — birinchi 2 yilda: {m['restaurants'][last]}+ restoran + "
             f"{m['families'][last]}+ uy xo'jaligi (moliyaviy modelga muvofiq)",
        "12": f"24/08/2002 — {age} yosh",
        "32": f"Oylik ~{mln(m['revenue'][0])} so'm ({m['restaurants'][0]} restoran + "
              f"{m['families'][0]} oila). Yillik ~{mln(m['revenue'][0] * 12)} so'm. "
              f"Mahsulotlar: mikro-ko'katlar, salatlar, gullar",
        "33": f"Jami: {m['customers'][0]} mijoz ({m['restaurants'][0]} restoran + "
              f"{m['families'][0]} oila). Sayt: microgreenuzbekistan.com — "
              f"50+ foydalanuvchi. Telegram-bot: 30+ obunachi",
        "35": plan,
        "10": f"${M.INVESTMENT_USD:,} — 2-siti-ferma (premium qulupnay gidroponik "
              f"ferma, 60 m2)".replace(",", " "),
    }
    for row in table.rows:
        number = row.cells[0].text.strip()
        if number in updates:
            set_cell(row.cells[2], updates[number])

    doc.save(OUT / NAMES["ariza"])


def main() -> int:
    m = M.build()
    OUT.mkdir(parents=True, exist_ok=True)
    build_chora(m)
    build_kpi(m)
    build_sotuv(m)
    build_biznes(m)
    build_ariza(m)
    for name in NAMES.values():
        print(f"[OK] {OUT / name}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
